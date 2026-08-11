"""文档解析层测试：txt/md/json/pdf/docx 解析、页码偏移与错误路径。

PDF/DOCX 测试文件在测试内动态生成（pypdf 写文本页 + python-docx 写段落表格），
不依赖仓库内的二进制 fixture。
"""
from __future__ import annotations

import io

import pytest

from mcp.document_parser import (
    SUPPORTED_EXTENSIONS,
    parse_document,
    _extract_docx,
    _extract_pdf,
)


# ── 测试内生成 PDF/DOCX 的辅助 ─────────────────────────────────────────────

def make_pdf(pages_text: list[str]) -> bytes:
    """用 pypdf 生成含文本页的 PDF（Helvetica 标准字体，纯 ASCII 文本）。"""
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    for text in pages_text:
        page = writer.add_blank_page(width=612, height=792)
        stream = StreamObject()
        safe = text.replace("\\", "\\\\").replace("(", r"\(").replace(")", r"\)")
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET".encode())
        page[NameObject("/Contents")] = stream
        page[NameObject("/Resources")] = DictionaryObject({
            NameObject("/Font"): DictionaryObject({
                NameObject("/F1"): DictionaryObject({
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                })
            })
        })
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """用 python-docx 生成含段落（可含表格）的 docx。"""
    import docx

    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        t = d.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                t.rows[i].cells[j].text = cell
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── txt / md ────────────────────────────────────────────────────────────────

def test_txt_whole_file_as_one_doc():
    docs = parse_document("新生指南.txt", "欢迎来到西电。".encode())
    assert len(docs) == 1
    assert docs[0]["title"] == "新生指南"
    assert docs[0]["content"] == "欢迎来到西电。"
    assert docs[0]["format"] == "txt"


def test_md_same_as_txt():
    docs = parse_document("排障手册.md", "# 标题\n正文内容".encode())
    assert docs[0]["title"] == "排障手册"
    assert docs[0]["format"] == "md"
    assert "# 标题\n正文内容" in docs[0]["content"]


def test_txt_empty_raises():
    with pytest.raises(ValueError, match="内容为空"):
        parse_document("空文档.txt", b"   \n ")


# ── json ────────────────────────────────────────────────────────────────────

def test_json_array_docs():
    data = '[{"title": "校历", "content": "秋季学期开学。"}, {"title": "选课", "content": "分预选正选。"}]'.encode()
    docs = parse_document("知识.json", data)
    assert len(docs) == 2
    assert docs[0]["title"] == "校历"
    assert docs[0]["format"] == "json"          # 自动补 format
    assert docs[1]["content"] == "分预选正选。"


def test_json_fields_passthrough():
    data = '[{"title": "校历", "content": "内容", "domain": "affairs", "source_url": "https://x"}]'.encode()
    doc = parse_document("知识.json", data)[0]
    assert doc["domain"] == "affairs"
    assert doc["source_url"] == "https://x"


def test_json_not_array_raises():
    with pytest.raises(ValueError, match="数组格式"):
        parse_document("坏.json", b'{"title": "x", "content": "y"}')


def test_json_broken_raises():
    with pytest.raises(ValueError, match="JSON 解析失败"):
        parse_document("坏.json", b'[{"title": ')


# ── pdf ─────────────────────────────────────────────────────────────────────

def test_pdf_multi_page_extracts_text_and_offsets():
    pdf = make_pdf(["Page one text. ", "Page two text. "])
    docs = parse_document("政策文件.pdf", pdf)
    assert len(docs) == 1
    doc = docs[0]
    assert doc["title"] == "政策文件"
    assert doc["format"] == "pdf"
    # 两页按 "\n\n" 拼接，偏移区间与页码对应
    full, offsets = doc["content"], doc["page_offsets"]
    assert full == "Page one text. \n\nPage two text. "
    assert len(offsets) == 2
    assert offsets[0] == (0, len("Page one text. "), 1)
    assert offsets[1][2] == 2
    assert offsets[1][0] == len("Page one text. ") + 2  # 跳过页分隔符


def test_pdf_scanned_no_text_layer_raises():
    pdf = make_pdf([" "])  # 空内容页 ≈ 扫描件无文本层
    with pytest.raises(ValueError, match="扫描件|无文本层"):
        parse_document("扫描件.pdf", pdf)


def test_pdf_corrupt_raises():
    with pytest.raises(ValueError, match="PDF"):
        parse_document("损坏.pdf", b"%PDF-1.4\nthis is not a real pdf file")


# ── docx ────────────────────────────────────────────────────────────────────

def test_docx_paragraphs_and_tables_in_order():
    docx = make_docx(
        ["第一条说明。", "第二条说明。"],
        table_rows=[["日期", "事项"], ["9月1日", "开学"]],
    )
    docs = parse_document("日程安排.docx", docx)
    assert len(docs) == 1
    text = docs[0]["content"]
    assert text.startswith("第一条说明。\n第二条说明。")
    assert "日期 | 事项" in text          # 表格行以 | 拼接，结构不丢
    assert "9月1日 | 开学" in text
    assert docs[0]["format"] == "docx"


def test_docx_empty_raises():
    empty = make_docx([])
    with pytest.raises(ValueError, match="空文档"):
        parse_document("空.docx", empty)


def test_extract_docx_skips_empty_lines():
    docx = make_docx(["段落一", "", "段落二"])
    assert _extract_docx(docx) == "段落一\n段落二"


# ── 其他 ────────────────────────────────────────────────────────────────────

def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="不支持的文件格式"):
        parse_document("表格.xlsx", b"whatever")


def test_no_extension_raises():
    with pytest.raises(ValueError, match="不支持的文件格式"):
        parse_document("README", b"text")


def test_supported_extensions():
    assert SUPPORTED_EXTENSIONS == {".txt", ".md", ".json", ".pdf", ".docx"}
